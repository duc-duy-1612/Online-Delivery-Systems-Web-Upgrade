import os
import pypandoc
from docx import Document
from docx.shared import Pt, Cm

md_file = r'c:\Users\ACER\Downloads\Online-Delivery-Systems-Web\docs\ba-portfolio\03_Requirements\User_Requirements_Document.md'
docx_file = r'c:\Users\ACER\Downloads\Online-Delivery-Systems-Web\docs\ba-portfolio\03_Requirements\User_Requirements_Document.docx'

# Download pandoc if not available
print("Downloading pandoc...")
pypandoc.download_pandoc()

print("Converting to DOCX...")
# Convert to DOCX
pypandoc.convert_file(
    md_file,
    'docx',
    outputfile=docx_file,
    extra_args=['-V', 'geometry:margin=2.5cm']
)

# Fix formatting in DOCX
doc = Document(docx_file)
# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.0)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

# Set font styles
for p in doc.paragraphs:
    if p.style.name.startswith('Heading'):
        for r in p.runs:
            r.font.name = 'Times New Roman'
    else:
        for r in p.runs:
            r.font.name = 'Times New Roman'
            if r.font.size is None:
                r.font.size = Pt(11)

for table in doc.tables:
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    if r.font.size is None:
                        r.font.size = Pt(10)

doc.save(docx_file)
print("DOCX Done!")
