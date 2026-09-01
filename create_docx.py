from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.0)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

# Set Default Font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 51) # dark navy

def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold:
        r.bold = True
    return p

# --- Content ---

add_heading("USER REQUIREMENTS DOCUMENT", 1)
add_heading("ONLINE FOOD DELIVERY SYSTEM", 2)

doc.add_paragraph("Document Version: 1.0")
doc.add_paragraph("Status: Baseline / Portfolio Evidence")
doc.add_paragraph("Prepared by: Pham Duc Duy")
doc.add_paragraph("Role: Business Analyst — Individual Portfolio Reconstruction & Validation")
doc.add_paragraph("Original Project: 2025")
doc.add_paragraph("BA Case Study Reconstruction: 2026")

doc.add_page_break()

# TOC (Simulated)
add_heading("Table of Contents", 1)
toc = [
    "1. Document Overview",
    "2. Business Context",
    "3. Stakeholders and User Groups",
    "4. User Personas / Actor Profiles",
    "5. User Goals",
    "6. User Requirements",
    "7. User Journeys and Cross-Role Handoffs",
    "8. User Constraints and Business Rules",
    "9. User Requirement Traceability",
    "10. Scope / Out of Scope",
    "11. Assumptions and Dependencies",
    "12. Open Clarifications",
    "13. Relationship to System / Functional Requirements",
    "14. Non-Functional User Expectations",
    "15. Glossary",
    "16. Document Relationship"
]
for item in toc:
    doc.add_paragraph(item)

doc.add_page_break()

add_heading("1. Document Overview", 1)
doc.add_paragraph("This User Requirements Document (URD) specifies the user-facing capabilities and outcomes required for the Online Food Delivery System. It is an individual portfolio reconstruction based on the approved project evidence baseline. The purpose of this URD is to document what the user groups need to accomplish and what outcomes they expect from the system, acting as a conceptual bridge between the high-level Business Context and the detailed Software and Functional Requirements Specifications.")
doc.add_paragraph("The URD focuses strictly on user needs, actor goals, and expected capabilities, rather than detailed system behavior and postconditions.")

add_heading("2. Business Context", 1)
doc.add_paragraph("The Online Food Delivery System operates in a multi-role food delivery marketplace environment. It facilitates food ordering and delivery fulfilment through structured interactions between Customers, Restaurants, and Shippers.")
doc.add_paragraph("- Customer: Browses menus, builds a cart, and completes checkout.", style='List Bullet')
doc.add_paragraph("- Restaurant: Receives incoming orders and prepares them for pickup.", style='List Bullet')
doc.add_paragraph("- Shipper: Accepts eligible delivery assignments and executes the physical delivery.", style='List Bullet')
doc.add_paragraph("- Administrator: Governs operational access and approvals for Restaurant and Shipper accounts.", style='List Bullet')
doc.add_paragraph("This digital platform replaces a disjointed manual process by centralizing operational data, enforcing controlled delivery-assignment eligibility, and improving end-to-end order lifecycle visibility.")

add_heading("3. Stakeholders and User Groups", 1)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Actor'
hdr_cells[1].text = 'Role in System'
hdr_cells[2].text = 'Primary Needs'

data = [
    ('Customer', 'Places and tracks orders', 'Browse, checkout, track, review'),
    ('Restaurant', 'Fulfils food orders', 'Process and prepare orders'),
    ('Shipper', 'Performs delivery', 'Accept, pickup, deliver, complete'),
    ('Administrator', 'Governs operational access', 'Approval and administration')
]
for item in data:
    row_cells = table.add_row().cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]
    row_cells[2].text = item[2]

add_heading("4. User Personas / Actor Profiles", 1)
add_heading("Customer", 2)
doc.add_paragraph("- Main objective: To discover food items, place orders, and receive them efficiently.", style='List Bullet')
doc.add_paragraph("- Primary interactions: Browsing menus, cart management, checkout, order tracking, and providing feedback.", style='List Bullet')
doc.add_paragraph("- Key information needs: Order status, delivery tracking, pricing/fee transparency.", style='List Bullet')
doc.add_paragraph("- Key constraints: Maximum delivery distance, valid payment selection.", style='List Bullet')

add_heading("Restaurant", 2)
doc.add_paragraph("- Main objective: To process incoming orders and prepare them for delivery.", style='List Bullet')
doc.add_paragraph("- Primary interactions: Menu management, accepting orders, updating order readiness.", style='List Bullet')
doc.add_paragraph("- Key information needs: Incoming order details, customer requests.", style='List Bullet')
doc.add_paragraph("- Key constraints: Requires Administrator approval for normal operations.", style='List Bullet')

add_heading("Shipper", 2)
doc.add_paragraph("- Main objective: To execute deliveries and earn income.", style='List Bullet')
doc.add_paragraph("- Primary interactions: Viewing available deliveries, accepting orders, updating delivery progress.", style='List Bullet')
doc.add_paragraph("- Key information needs: Pickup location, drop-off location, order readiness status.", style='List Bullet')
doc.add_paragraph("- Key constraints: Can only accept unassigned orders marked ready; limited to one active delivery; requires Administrator approval.", style='List Bullet')

add_heading("Administrator", 2)
doc.add_paragraph("- Main objective: To manage and govern operational participants.", style='List Bullet')
doc.add_paragraph("- Primary interactions: Approving Partner registrations, viewing system statistics.", style='List Bullet')
doc.add_paragraph("- Key information needs: Partner registration details, platform performance statistics.", style='List Bullet')
doc.add_paragraph("- Key constraints: Governs but does not participate in the core order lifecycle.", style='List Bullet')

add_heading("5. User Goals", 1)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Actor'
hdr_cells[1].text = 'User Goal ID'
hdr_cells[2].text = 'User Goal'

data = [
    ('Customer', 'UG-CUS-01', 'Discover restaurants and products'),
    ('Customer', 'UG-CUS-02', 'Build a shopping cart'),
    ('Customer', 'UG-CUS-03', 'Complete an eligible checkout'),
    ('Customer', 'UG-CUS-04', 'Track order status and delivery progress'),
    ('Customer', 'UG-CUS-05', 'Submit a post-completion review'),
    ('Restaurant', 'UG-RES-01', 'Access approved operational capabilities'),
    ('Restaurant', 'UG-RES-02', 'Process and prepare incoming orders'),
    ('Restaurant', 'UG-RES-03', 'Indicate order readiness for pickup'),
    ('Shipper', 'UG-SHP-01', 'Identify eligible delivery assignments'),
    ('Shipper', 'UG-SHP-02', 'Accept an eligible order'),
    ('Shipper', 'UG-SHP-03', 'Execute pickup and update delivery progress'),
    ('Shipper', 'UG-SHP-04', 'Complete delivery'),
    ('Administrator', 'UG-ADM-01', 'Manage user accounts and approve mapped operational participants'),
    ('Administrator', 'UG-ADM-02', 'Maintain administrative control and view statistics')
]
for item in data:
    row_cells = table.add_row().cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]
    row_cells[2].text = item[2]

add_heading("6. User Requirements", 1)

def add_ur(ur_id, title, need, outcome, related):
    p = doc.add_paragraph()
    p.add_run(ur_id + "\n").bold = True
    p.add_run("Title: ").bold = True
    p.add_run(title + "\n")
    p.add_run("User Need: ").bold = True
    p.add_run(need + "\n")
    p.add_run("Expected Outcome: ").bold = True
    p.add_run(outcome + "\n")
    p.add_run("Related Requirements: ").bold = True
    p.add_run(related)

add_heading("Customer User Requirements", 2)
add_ur("UR-CUS-01", "Account & Access", "The Customer needs to register, access the system using the supported authentication mechanism, and manage their personal profile.", "The Customer can access permitted account functions and profile settings after successful authentication.", "FR-AUTH-01, FR-AUTH-02, FR-CUS-01")
add_ur("UR-CUS-02", "Restaurant / Product Discovery", "The Customer needs to browse active restaurants and view their associated menus.", "The Customer can discover available food items to order.", "FR-CUS-02")
add_ur("UR-CUS-03", "Shopping Cart", "The Customer needs to add/remove items and adjust quantities in a shopping cart.", "The Customer can prepare their desired order and see cart totals before proceeding to checkout.", "FR-CUS-03")
add_ur("UR-CUS-04", "Checkout", "The Customer needs to provide delivery information, validate the address/distance, review calculated fees (delivery and service), select a payment method (COD or QR Payment Simulation), and submit an eligible order.", "The Customer successfully submits an order if delivery conditions are met, initiating the order lifecycle.", "FR-CUS-04, FR-CUS-05, FR-CUS-06")
add_ur("UR-CUS-05", "Order Tracking", "The Customer needs to view current order status and monitor delivery progress/route information.", "The Customer is informed of the order's progress throughout the fulfilment lifecycle.", "FR-CUS-07")
add_ur("UR-CUS-06", "History", "The Customer needs to view their current orders and completed order history.", "The Customer can access details of past and present orders.", "FR-CUS-09")
add_ur("UR-CUS-07", "Review", "The Customer needs to submit ratings and reviews for the Restaurant and Shipper after an eligible order is completed.", "The Customer can provide feedback only after the order reaches completion.", "FR-CUS-08")

add_heading("Restaurant User Requirements", 2)
add_ur("UR-RES-01", "Profile & Menu Management", "The Restaurant needs to access approved operational capabilities to update store information, manage menu categories, and manage food items.", "The Restaurant can maintain an accurate profile and current menu offerings for Customer discovery.", "FR-RES-01, FR-RES-02, FR-RES-03")
add_ur("UR-RES-02", "Order Processing", "The Restaurant needs to receive/view incoming mapped order information, prepare the order, and indicate when the food is ready for pickup.", "The Restaurant can process orders and trigger the readiness update to make the order available for Shippers.", "FR-RES-04, FR-RES-05")
add_ur("UR-RES-03", "History & Reviews", "The Restaurant needs to view revenue, order history, and customer reviews.", "The Restaurant can track business performance and monitor customer feedback.", "FR-RES-06, FR-RES-07")

add_heading("Shipper User Requirements", 2)
add_ur("UR-SHP-01", "Delivery Assignment", "The Shipper needs to access unassigned, eligible delivery opportunities and accept an eligible order.", "The Shipper is officially assigned to execute a delivery, provided the order is unassigned, ready for pickup, and the Shipper meets active delivery constraints.", "FR-SHP-01, FR-SHP-02")
add_ur("UR-SHP-02", "Delivery Execution", "The Shipper needs to perform the pickup, update delivery progress, and complete the delivery within the system.", "The Order lifecycle correctly advances to the delivery and completion states based on Shipper updates.", "FR-SHP-03")
add_ur("UR-SHP-03", "Profile, History & Income", "The Shipper needs to manage their profile and view their delivery history and accumulated income.", "The Shipper can monitor their earnings and maintain up-to-date account details.", "FR-SHP-04, FR-SHP-05")

add_heading("Administrator User Requirements", 2)
add_ur("UR-ADM-01", "Partner Approval", "The Administrator needs to approve Restaurants and Shippers for mapped normal operational capabilities.", "Only approved partner accounts are granted access to core order-processing operations.", "FR-ADM-02")
add_ur("UR-ADM-02", "User & System Administration", "The Administrator needs to manage user accounts, view system statistics, and export revenue statistics.", "The Administrator maintains governance control and can report on platform operations.", "FR-ADM-01, FR-ADM-03, FR-ADM-04")

add_heading("7. User Journeys and Cross-Role Handoffs", 1)
add_heading("End-to-End User Journey", 2)
doc.add_paragraph("Customer Browse\n↓\nSelect Products\n↓\nCart\n↓\nCheckout\n↓\nAddress / Distance Validation\n↓\nFee Calculation\n↓\nCreate Order\n↓\nChờ xác nhận\n↓\nRestaurant Preparation\n↓\nĐang lấy món\n↓\nShipper Pickup\n↓\nĐang giao\n↓\nDelivery Completion\n↓\nHoàn thành\n↓\nCustomer Review")

add_heading("Cross-Role Handoffs", 2)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'From'
hdr_cells[1].text = 'To'
hdr_cells[2].text = 'Trigger / Handoff'
hdr_cells[3].text = 'User Need'

data = [
    ('Customer', 'System', 'Checkout submission', 'Customer needs to place the order'),
    ('System', 'Restaurant', 'Order created / confirmed', 'Restaurant needs to receive order details'),
    ('Restaurant', 'System', 'Order ready', 'Restaurant needs to indicate preparation is complete'),
    ('System', 'Shipper', 'Delivery Opportunity', 'Shipper needs to view unassigned, ready orders'),
    ('Shipper', 'System', 'Delivery acceptance', 'Shipper needs to accept the assignment'),
    ('Shipper', 'Customer', 'Delivery progression', 'Customer needs to monitor delivery progress'),
    ('Shipper', 'System', 'Delivery completion', 'Shipper needs to complete the delivery')
]
for item in data:
    row_cells = table.add_row().cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]
    row_cells[2].text = item[2]
    row_cells[3].text = item[3]

add_heading("8. User Constraints and Business Rules", 1)
doc.add_paragraph("The following constraints are supported by the current baseline and govern the user requirements:")
doc.add_paragraph("- Maximum delivery route distance: Checkout is constrained to a 30 km delivery radius.", style='List Bullet')
doc.add_paragraph("- Fee calculation rules: Distance-based delivery fees and time-based service fees are applied during checkout.", style='List Bullet')
doc.add_paragraph("- Shipper Assignment constraints: An order must be unassigned and in \"Đang lấy món\" for a Shipper to accept it.", style='List Bullet')
doc.add_paragraph("- Active Delivery limit: A Shipper cannot hold multiple active delivery orders simultaneously.", style='List Bullet')
doc.add_paragraph("- Completion requirement: An order must reach the \"Hoàn thành\" state before a Customer is eligible to submit a review.", style='List Bullet')
doc.add_paragraph("- Partner Approval requirement: Restaurant and Shipper accounts require Administrator approval before engaging in normal operations.", style='List Bullet')

add_heading("9. User Requirement Traceability", 1)
table = doc.add_table(rows=1, cols=8)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
headers = ['URD ID', 'Actor', 'User Requirement', 'Related FR', 'Related UC', 'Related US', 'Related AC', 'Related BR']
for i, h in enumerate(headers):
    hdr_cells[i].text = h

data = [
    ('UR-CUS-01', 'Customer', 'Account & Access', 'FR-AUTH-01, FR-AUTH-02, FR-CUS-01', 'UC-CUS-01', 'US-CUS-01', 'AC-US-CUS-01-01 to 03', 'Not directly mapped in current baseline'),
    ('UR-CUS-02', 'Customer', 'Restaurant / Product Discovery', 'FR-CUS-02', 'UC-CUS-02', 'US-CUS-02', 'AC-US-CUS-02-01', 'Not directly mapped in current baseline'),
    ('UR-CUS-03', 'Customer', 'Shopping Cart', 'FR-CUS-03', 'UC-CUS-02', 'US-CUS-03', 'AC-US-CUS-03-01 to 03', 'Not directly mapped in current baseline'),
    ('UR-CUS-04', 'Customer', 'Checkout', 'FR-CUS-04, FR-CUS-05, FR-CUS-06', 'UC-CUS-03', 'US-CUS-04', 'AC-US-CUS-04-01 to 09', 'BR-DEL-01, BR-FEE-01, BR-FEE-02'),
    ('UR-CUS-05', 'Customer', 'Order Tracking', 'FR-CUS-07', 'UC-CUS-04', 'US-CUS-05', 'AC-US-CUS-05-01; 02', 'Not directly mapped in current baseline'),
    ('UR-CUS-06', 'Customer', 'History', 'FR-CUS-09', 'UC-CUS-05', 'US-CUS-06', 'AC-US-CUS-06-01 to 03', 'Not directly mapped in current baseline'),
    ('UR-CUS-07', 'Customer', 'Review', 'FR-CUS-08', 'UC-CUS-06', 'US-CUS-07', 'AC-US-CUS-07-01; 02', 'BR-ORDER-01'),
    ('UR-RES-01', 'Restaurant', 'Profile & Menu Management', 'FR-RES-01, FR-RES-02, FR-RES-03', 'UC-RES-01', 'US-RES-01, 02, 03', 'AC-US-RES-01-01, 02-01, 03-01 to 03', 'Not directly mapped in current baseline'),
    ('UR-RES-02', 'Restaurant', 'Order Processing', 'FR-RES-04, FR-RES-05', 'UC-RES-02', 'US-RES-04', 'AC-US-RES-04-01; 02', 'BR-PARTNER-01'),
    ('UR-RES-03', 'Restaurant', 'History & Reviews', 'FR-RES-06, FR-RES-07', 'UC-RES-03', 'US-RES-05, 06', 'AC-US-RES-05-01; 02, 06-01', 'Not directly mapped in current baseline'),
    ('UR-SHP-01', 'Shipper', 'Delivery Assignment', 'FR-SHP-01, FR-SHP-02', 'UC-SHP-01, 02', 'US-SHP-01, 02', 'AC-US-SHP-01-01, 02-01 to 04', 'BR-SHIP-01, BR-SHIP-02, BR-PARTNER-01'),
    ('UR-SHP-02', 'Shipper', 'Delivery Execution', 'FR-SHP-03', 'UC-SHP-03', 'US-SHP-03', 'AC-US-SHP-03-01 to 03', 'BR-PARTNER-01'),
    ('UR-SHP-03', 'Shipper', 'Profile, History & Income', 'FR-SHP-04, FR-SHP-05', 'UC-SHP-04, 05', 'US-SHP-04, 05', 'AC-US-SHP-04-01; 02, 05-01; 02', 'Not directly mapped in current baseline'),
    ('UR-ADM-01', 'Administrator', 'Partner Approval', 'FR-ADM-02', 'UC-ADM-02', 'US-ADM-02', 'AC-US-ADM-02-01', 'BR-PARTNER-01'),
    ('UR-ADM-02', 'Administrator', 'User & System Administration', 'FR-ADM-01, FR-ADM-03, FR-ADM-04', 'UC-ADM-01, 03', 'US-ADM-01, 03, 04', 'AC-US-ADM-01-01, 03-01; 02, 04-01', 'Not directly mapped in current baseline')
]
for item in data:
    row_cells = table.add_row().cells
    for i in range(8):
        row_cells[i].text = item[i]

add_heading("10. Scope / Out of Scope", 1)
add_paragraph("In Scope:", bold=True)
doc.add_paragraph("- Core order and delivery lifecycle involving the Customer, Restaurant, Shipper, and Administrator.", style='List Bullet')
doc.add_paragraph("- Responsive Web Application implementation.", style='List Bullet')
doc.add_paragraph("- Supported order state transitions: Chờ xác nhận, Đang lấy món, Đang giao, Hoàn thành.", style='List Bullet')
add_paragraph("Out of Scope:", bold=True)
doc.add_paragraph("- Native mobile applications.", style='List Bullet')
doc.add_paragraph("- Real production payment-gateway integration.", style='List Bullet')

add_heading("11. Assumptions and Dependencies", 1)
add_paragraph("Assumptions:", bold=True)
doc.add_paragraph("- Restaurant and Shipper normal operational capabilities require Administrator approval as defined by BR-PARTNER-01.", style='List Bullet')
doc.add_paragraph("- COD / QR is represented as payment simulation rather than production payment-gateway integration.", style='List Bullet')
add_paragraph("Dependencies:", bold=True)
doc.add_paragraph("- Administrator approval is a prerequisite for partner operational flows.", style='List Bullet')
doc.add_paragraph("- Address / route-distance calculation functionality must be supported by Map APIs.", style='List Bullet')
doc.add_paragraph("- Order-state transitions dictate downstream participant eligibility (e.g., Shipper eligibility relies on Restaurant readiness).", style='List Bullet')

add_heading("12. Open Clarifications", 1)
doc.add_paragraph("The following semantics remain explicitly unresolved in the current baseline and affect user requirements:")
doc.add_paragraph("- Definition of \"active delivery\" for the single active delivery constraint.", style='List Bullet')
doc.add_paragraph("- Fractional-kilometre rounding logic for distance-based fee calculation.", style='List Bullet')
doc.add_paragraph("- Authoritative timestamp for the service-fee rule.", style='List Bullet')
doc.add_paragraph("- Wider non-core Partner capability boundary regarding Administrator approval.", style='List Bullet')
doc.add_paragraph("- Exact post-Shipper-cancellation Order.Status transition.", style='List Bullet')
doc.add_paragraph("- Exact Restaurant confirmation semantics between \"Chờ xác nhận\" and \"Đang lấy món\".", style='List Bullet')
doc.add_paragraph("- Review cardinality and editability rules.", style='List Bullet')

add_heading("13. Relationship to System / Functional Requirements", 1)
doc.add_paragraph("This URD describes the \"What\" and \"Why\" from the user's perspective, providing context to the Functional Requirements Specification (FRS) which describes the \"How\" from the system's perspective.")
doc.add_paragraph("Example translation:\n\nUR-CUS-04\nCustomer needs validated checkout.\n\n↓ translated into\n\nFR-CUS-04\nValidate Address & Distance\n\n↓ supported by\n\nUC-CUS-03\nCheckout and Place Order\n\n↓ expressed as\n\nUS-CUS-04\nComplete Checkout\n\n↓ tested by\n\nAC-US-CUS-04-02 / 03")
doc.add_paragraph("This ensures requirements are user-centered and accurately traced through the software delivery lifecycle.")

add_heading("14. Non-Functional User Expectations", 1)
doc.add_paragraph("Non-functional user expectations are not fully defined in the current baseline and are outside the confirmed scope of this URD. The system relies on the approved NFRs documented in the SRS.")

add_heading("15. Glossary", 1)
doc.add_paragraph("- Administrator: Governs user accounts and operational approval.", style='List Bullet')
doc.add_paragraph("- COD: Cash on Delivery.", style='List Bullet')
doc.add_paragraph("- Customer: The user discovering items and placing orders.", style='List Bullet')
doc.add_paragraph("- Order.Status: The formal lifecycle state of an order (Chờ xác nhận, Đang lấy món, Đang giao, Hoàn thành).", style='List Bullet')
doc.add_paragraph("- QR Payment Simulation: Simulated payment mechanism for non-cash checkout.", style='List Bullet')
doc.add_paragraph("- Restaurant: The partner processing and preparing food orders.", style='List Bullet')
doc.add_paragraph("- RTM: Requirements Traceability Matrix.", style='List Bullet')
doc.add_paragraph("- Shipper: The partner accepting and executing physical delivery.", style='List Bullet')
doc.add_paragraph("- UAT: User Acceptance Testing.", style='List Bullet')
doc.add_paragraph("- TARGET: The intended future state capabilities.", style='List Bullet')
doc.add_paragraph("- CURRENT: The existing implementation behavior.", style='List Bullet')

add_heading("16. Document Relationship", 1)
doc.add_paragraph("BRD\n ↓\nURD\n ↓\nSRS\n ↓\nFRS\n ↓\nUC / US / AC\n ↓\nBusiness Rules\n ↓\nState / ERD / Data Dictionary\n ↓\nRTM\n ↓\nUAT")
doc.add_paragraph("(This represents the controlled documentation relationship, not necessarily strict historical creation order.)")

# Apply font to everything in tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    if r.font.size is None:
                        r.font.size = Pt(10)

doc.save(r'c:\Users\ACER\Downloads\Online-Delivery-Systems-Web\docs\ba-portfolio\03_Requirements\User_Requirements_Document.docx')
print("DOCX created successfully.")
