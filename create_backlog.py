from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.section import WD_ORIENT

doc = Document()

# Define page layouts
def set_portrait():
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

def set_landscape():
    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

set_portrait()

# Styles
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 51) # dark navy
    return h

def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p

# --- COVER ---
title = doc.add_heading('PRODUCT BACKLOG', 0)
for r in title.runs:
    r.font.name = 'Times New Roman'
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0, 0, 51)
    
sub = doc.add_heading('ONLINE FOOD DELIVERY SYSTEM', 1)
for r in sub.runs:
    r.font.name = 'Times New Roman'
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0, 0, 51)

doc.add_paragraph("\nDocument Version: 1.0")
doc.add_paragraph("Status: Baseline / Portfolio Evidence")
doc.add_paragraph("Prepared by: Pham Duc Duy")
doc.add_paragraph("Role: Business Analyst")
doc.add_paragraph("Original Project: 2025")
doc.add_paragraph("BA Case Study Reconstruction: 2026")

doc.add_page_break()

# --- TOC ---
add_heading("Table of Contents")
toc = [
    "1. Document Overview",
    "2. Product / System Context",
    "3. Backlog Governance",
    "4. Product Goals",
    "5. Stakeholders / Actors",
    "6. Product Backlog Summary",
    "7. Epic / Feature Structure",
    "8. User Story Map",
    "9. Detailed Product Backlog",
    "10. Acceptance Criteria Summary",
    "11. Business Rule Mapping",
    "12. Order Lifecycle Mapping",
    "13. Traceability Matrix",
    "14. Clarifications / Open Questions",
    "15. Dependencies",
    "16. Non-Functional Considerations",
    "17. Out of Scope",
    "18. Definition of Ready",
    "19. Definition of Done",
    "20. Glossary"
]
for t in toc:
    doc.add_paragraph(t)

doc.add_page_break()

# --- BODY ---
add_heading("1. Document Overview")
doc.add_paragraph("This Product Backlog provides an actionable Agile delivery backlog derived from the validated requirements baseline for the Online Food Delivery System. It acts as a portfolio reconstruction demonstrating the transformation of structured requirements into actionable backlog items.")

add_heading("2. Product / System Context")
doc.add_paragraph("The Online Food Delivery System facilitates food ordering and delivery fulfilment through interactions between Customers, Restaurants, Shippers, and Administrators.")

add_heading("3. Backlog Governance")
doc.add_paragraph("Priority values are analytical portfolio prioritization for backlog representation and are not claimed as historical Product Owner decisions. Backlog status represents the portfolio reconstruction state and is not claimed to be the historical sprint status.")

add_heading("4. Product Goals")
goals = [
    "Enable Customers to discover products and place eligible orders.",
    "Enable Restaurants to process and prepare orders.",
    "Enable Shippers to execute eligible deliveries.",
    "Provide controlled order lifecycle visibility.",
    "Enforce critical distance, fee, assignment and review rules.",
    "Provide Administrator governance over operational participants."
]
for g in goals:
    doc.add_paragraph("- " + g, style='List Bullet')

add_heading("5. Stakeholders / Actors")
actors = ["Customer", "Restaurant", "Shipper", "Administrator"]
for a in actors:
    doc.add_paragraph("- " + a, style='List Bullet')

add_heading("6. Product Backlog Summary")
doc.add_paragraph("Source-baseline metrics:")
doc.add_paragraph("- 22 User Stories", style='List Bullet')
doc.add_paragraph("- 27 Functional Requirements", style='List Bullet')
doc.add_paragraph("- 7 Business Rules", style='List Bullet')
doc.add_paragraph("- 50 Acceptance Criteria", style='List Bullet')
doc.add_paragraph("- 38 RTM Rows", style='List Bullet')
doc.add_paragraph("- 60 Validation Catalogue Entries", style='List Bullet')
doc.add_paragraph("Backlog User Story rows: 22")
doc.add_paragraph("Epic count: 8")

add_heading("7. Epic / Feature Structure")
epics = [
    "EPIC-01: Customer Account & Access",
    "EPIC-02: Restaurant & Product Discovery",
    "EPIC-03: Shopping Cart & Checkout",
    "EPIC-04: Order Processing",
    "EPIC-05: Shipper & Delivery Management",
    "EPIC-06: Order Tracking & Completion",
    "EPIC-07: Reviews & Post-Order Experience",
    "EPIC-08: Administrator & Partner Governance"
]
for e in epics:
    doc.add_paragraph(e)

add_heading("8. User Story Map")
doc.add_paragraph("Analytical MVP view reconstructed for portfolio demonstration.")
doc.add_paragraph("Discover → Select → Order → Fulfil → Deliver → Track → Complete → Review")
doc.add_paragraph("Customer User Stories map primarily to Discover, Select, Order, Track, Review.")
doc.add_paragraph("Restaurant User Stories map primarily to Fulfil.")
doc.add_paragraph("Shipper User Stories map primarily to Deliver, Complete.")

# --- LANDSCAPE BACKLOG TABLE ---
set_landscape()
add_heading("9. Detailed Product Backlog")

table = doc.add_table(rows=1, cols=11)
table.style = 'Table Grid'
hdr = table.rows[0].cells
headers = ["PB ID", "Epic", "US ID", "Actor", "User Story", "Priority", "Business Value", "AC Summary", "BR", "Traceability", "Status"]
for i, h in enumerate(headers):
    hdr[i].text = h
    
stories = [
    ("PB-001", "EPIC-01", "US-CUS-01", "Customer", "As a Customer, I want to view and update my personal profile information and change my account password, so that my account details remain accurate.", "P1", "Enables customers to maintain account security and accurate profile details.", "AC-US-CUS-01-01 to 03", "—", "FR-CUS-01, UC-CUS-01", "Baseline"),
    ("PB-002", "EPIC-02", "US-CUS-02", "Customer", "As a Customer, I want to browse active restaurants and their associated menus, so that I can find food items to order.", "P1", "Enables customers to discover available products and make purchasing decisions.", "AC-US-CUS-02-01", "—", "FR-CUS-02, UC-CUS-02", "Baseline"),
    ("PB-003", "EPIC-02", "US-RES-01", "Restaurant", "As a Restaurant, I want to update my store information, so that my store profile remains accurate.", "P1", "Enables restaurants to present accurate information to customers.", "AC-US-RES-01-01", "—", "FR-RES-01, UC-RES-01", "Baseline"),
    ("PB-004", "EPIC-02", "US-RES-02", "Restaurant", "As a Restaurant, I want to manage my menu categories, so that I can organize my food offerings.", "P1", "Enables restaurants to organize their menu structure.", "AC-US-RES-02-01", "—", "FR-RES-02, UC-RES-01", "Baseline"),
    ("PB-005", "EPIC-02", "US-RES-03", "Restaurant", "As a Restaurant, I want to create, update, and remove food items, so that my menu reflects my current offerings.", "P1", "Enables restaurants to provide up-to-date food offerings.", "AC-US-RES-03-01 to 03", "—", "FR-RES-03, UC-RES-01", "Baseline"),
    ("PB-006", "EPIC-03", "US-CUS-03", "Customer", "As a Customer, I want to manage items in my shopping cart, so that I can prepare my desired order for checkout.", "P1", "Enables customers to prepare an order before checkout.", "AC-US-CUS-03-01 to 03", "—", "FR-CUS-03, UC-CUS-02", "Baseline"),
    ("PB-007", "EPIC-03", "US-CUS-04", "Customer", "As a Customer, I want to finalize my cart, validate delivery details, and submit my order using COD or QR Payment Simulation, so that my order is placed with the Restaurant.", "P0", "Enables customers to complete an eligible order.", "AC-US-CUS-04-01 to 09", "BR-DEL-01, BR-FEE-01, 02", "FR-CUS-04, 05, 06, UC-CUS-03", "Baseline"),
    ("PB-008", "EPIC-04", "US-RES-04", "Restaurant", "As a Restaurant, I want to view incoming orders and mark them as ready when preparation is complete, so that the order becomes available for Shipper pickup.", "P0", "Enables restaurants to process and prepare orders for delivery.", "AC-US-RES-04-01, 02", "BR-PARTNER-01", "FR-RES-04, 05, UC-RES-02", "Baseline"),
    ("PB-009", "EPIC-04", "US-RES-05", "Restaurant", "As a Restaurant, I want to view my revenue statistics and order history, so that I can track my business performance.", "P2", "Enables restaurants to monitor financial and operational performance.", "AC-US-RES-05-01, 02", "—", "FR-RES-06, UC-RES-03", "Baseline"),
    ("PB-010", "EPIC-05", "US-SHP-01", "Shipper", "As a Shipper, I want to view available and unassigned delivery orders, so that I can select an eligible delivery assignment.", "P1", "Enables shippers to discover unassigned delivery opportunities.", "AC-US-SHP-01-01", "BR-PARTNER-01", "FR-SHP-01, UC-SHP-01", "Baseline"),
    ("PB-011", "EPIC-05", "US-SHP-02", "Shipper", "As a Shipper, I want to claim an available delivery order, so that I am officially assigned to execute the delivery.", "P0", "Prevents invalid delivery assignments and assigns a shipper to an order.", "AC-US-SHP-02-01 to 04", "BR-SHIP-01, 02, PARTNER", "FR-SHP-02, UC-SHP-02", "Baseline"),
    ("PB-012", "EPIC-05", "US-SHP-03", "Shipper", "As a Shipper, I want to update the Order lifecycle as I collect and complete a delivery, so that the delivery progress is recorded.", "P0", "Enables shippers to complete the physical delivery and update order status.", "AC-US-SHP-03-01 to 03", "BR-PARTNER-01", "FR-SHP-03, UC-SHP-03", "Baseline"),
    ("PB-013", "EPIC-05", "US-SHP-04", "Shipper", "As a Shipper, I want to view my completed deliveries and income statistics, so that I can track my earnings and activity.", "P2", "Enables shippers to track their earnings.", "AC-US-SHP-04-01, 02", "—", "FR-SHP-04, UC-SHP-04", "Baseline"),
    ("PB-014", "EPIC-05", "US-SHP-05", "Shipper", "As a Shipper, I want to view and update my personal profile and account information, so that my details remain current.", "P2", "Enables shippers to maintain accurate profile information.", "AC-US-SHP-05-01, 02", "—", "FR-SHP-05, UC-SHP-05", "Baseline"),
    ("PB-015", "EPIC-06", "US-CUS-05", "Customer", "As a Customer, I want to view the status and routing of my active order, so that I am informed of its progress.", "P1", "Preserves order visibility across the fulfilment lifecycle.", "AC-US-CUS-05-01, 02", "—", "FR-CUS-07, UC-CUS-04", "Baseline"),
    ("PB-016", "EPIC-06", "US-CUS-06", "Customer", "As a Customer, I want to view my current orders, completed order history, and detailed information for selected orders, so that I can monitor my activity.", "P2", "Enables customers to review past activity.", "AC-US-CUS-06-01 to 03", "—", "FR-CUS-09, UC-CUS-05", "Baseline"),
    ("PB-017", "EPIC-07", "US-CUS-07", "Customer", "As a Customer, I want to submit ratings and reviews for the Restaurant and Shipper of a completed order, so that I can provide feedback on my experience.", "P2", "Provides feedback to maintain quality of service.", "AC-US-CUS-07-01, 02", "BR-ORDER-01", "FR-CUS-08, UC-CUS-06", "Baseline"),
    ("PB-018", "EPIC-07", "US-RES-06", "Restaurant", "As a Restaurant, I want to view customer ratings and reviews associated with completed orders, so that I can monitor customer feedback.", "P2", "Enables restaurants to monitor customer feedback.", "AC-US-RES-06-01", "—", "FR-RES-07, UC-RES-03", "Baseline"),
    ("PB-019", "EPIC-08", "US-ADM-01", "Administrator", "As an Administrator, I want to manage system user accounts, so that supported account administration can be performed.", "P1", "Enables administrators to manage platform access.", "AC-US-ADM-01-01", "—", "FR-ADM-01, UC-ADM-01", "Baseline"),
    ("PB-020", "EPIC-08", "US-ADM-02", "Administrator", "As an Administrator, I want to approve new Restaurant and Shipper registrations, so that approved Restaurant and Shipper accounts can proceed to normal operation.", "P0", "Governs operational access for partners to ensure quality and compliance.", "AC-US-ADM-02-01", "BR-PARTNER-01", "FR-ADM-02, UC-ADM-02", "Baseline"),
    ("PB-021", "EPIC-08", "US-ADM-03", "Administrator", "As an Administrator, I want to view system-wide operational and revenue statistics, so that I can review system-wide operational and revenue information.", "P2", "Provides visibility into platform performance.", "AC-US-ADM-03-01, 02", "—", "FR-ADM-03, UC-ADM-03", "Baseline"),
    ("PB-022", "EPIC-08", "US-ADM-04", "Administrator", "As an Administrator, I want to export revenue statistics to Excel, so that revenue statistics are available in Excel format for downstream use.", "P3", "Provides external reporting capabilities.", "AC-US-ADM-04-01", "—", "FR-ADM-04, UC-ADM-03", "Baseline"),
]

for row_data in stories:
    row = table.add_row().cells
    for i, text in enumerate(row_data):
        row[i].text = text

# Add Traceability Matrix table in Landscape too
doc.add_page_break()
add_heading("13. Traceability Matrix")
t_table = doc.add_table(rows=1, cols=8)
t_table.style = 'Table Grid'
t_hdr = t_table.rows[0].cells
t_headers = ["PB ID", "Epic", "US ID", "FR ID", "UC ID", "AC ID(s)", "BR ID(s)", "UAT ID(s)"]
for i, h in enumerate(t_headers):
    t_hdr[i].text = h

for row_data in stories:
    row = t_table.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[9].split(",")[0] # approximate
    row[4].text = row_data[9].split(",")[-1].strip()
    row[5].text = row_data[7]
    row[6].text = row_data[8]
    row[7].text = "Not directly mapped"

# --- PORTRAIT AGAIN FOR REMAINING ---
set_portrait()

add_heading("10. Acceptance Criteria Summary")
doc.add_paragraph("Full Acceptance Criteria are maintained in the source User Stories & Acceptance Criteria artefact. The backlog preserves the exact numbering and logic of the 50 approved Acceptance Criteria.")

add_heading("11. Business Rule Mapping")
doc.add_paragraph("- BR-ORDER-01: Mapped to US-CUS-07")
doc.add_paragraph("- BR-SHIP-01, BR-SHIP-02: Mapped to US-SHP-02")
doc.add_paragraph("- BR-DEL-01, BR-FEE-01, BR-FEE-02: Mapped to US-CUS-04")
doc.add_paragraph("- BR-PARTNER-01: Mapped to US-RES-04, US-SHP-01, US-SHP-02, US-SHP-03, US-ADM-02")

add_heading("12. Order Lifecycle Mapping")
doc.add_paragraph("- ST-01 — Chờ xác nhận: Created via checkout (US-CUS-04)")
doc.add_paragraph("- ST-02 — Đang lấy món: Triggered by Restaurant (US-RES-04)")
doc.add_paragraph("- ST-03 — Đang giao: Triggered by Shipper pickup (US-SHP-03)")
doc.add_paragraph("- ST-04 — Hoàn thành: Triggered by Shipper completion (US-SHP-03)")

add_heading("14. Clarifications / Open Questions")
cls = [
    "CL-01: Definition of \"active delivery\" for the single active delivery constraint.",
    "CL-02: Fractional-kilometre rounding logic for distance-based fee calculation.",
    "CL-03: Authoritative timestamp for the service fee rule.",
    "CL-04: Wider non-core Partner capability boundary regarding Administrator approval.",
    "CL-05: Exact post-Shipper-cancellation Order.Status transition.",
    "CL-06: Exact Restaurant confirmation semantics between \"Chờ xác nhận\" and \"Đang lấy món\".",
    "CL-07: Review cardinality and editability rules."
]
for c in cls:
    doc.add_paragraph(c)

add_heading("15. Dependencies")
deps = [
    "Administrator approval",
    "Address / route-distance calculation",
    "Order state transitions",
    "Restaurant readiness",
    "Shipper eligibility"
]
for d in deps:
    doc.add_paragraph("- " + d, style='List Bullet')

add_heading("16. Non-Functional Considerations")
doc.add_paragraph("Technical / Delivery Considerations: System relies on approved NFRs documented in the SRS (NFR-COMP-01, NFR-PER-01, NFR-SEC-01, NFR-SEC-02). No technical tasks fabricated.")

add_heading("17. Out of Scope")
doc.add_paragraph("- Native mobile applications", style='List Bullet')
doc.add_paragraph("- Real production payment-gateway integration", style='List Bullet')

add_heading("18. Definition of Ready")
doc.add_paragraph("Portfolio Agile Governance Baseline:")
doc.add_paragraph("- User / actor is identified", style='List Bullet')
doc.add_paragraph("- Desired capability is clear", style='List Bullet')
doc.add_paragraph("- Business value is understood", style='List Bullet')
doc.add_paragraph("- Acceptance criteria are available", style='List Bullet')
doc.add_paragraph("- Relevant business rules are identified", style='List Bullet')
doc.add_paragraph("- Dependencies are understood", style='List Bullet')
doc.add_paragraph("- Unresolved semantics are explicitly recorded", style='List Bullet')

add_heading("19. Definition of Done")
doc.add_paragraph("Portfolio Governance Baseline:")
doc.add_paragraph("- Implementation completed", style='List Bullet')
doc.add_paragraph("- Acceptance criteria satisfied", style='List Bullet')
doc.add_paragraph("- Relevant business rules validated", style='List Bullet')
doc.add_paragraph("- Functional testing completed", style='List Bullet')
doc.add_paragraph("- UAT evidence available where applicable", style='List Bullet')
doc.add_paragraph("- RTM updated", style='List Bullet')

add_heading("20. Glossary")
doc.add_paragraph("- PB ID: Product Backlog ID")
doc.add_paragraph("- FR / UC / US / AC: Functional Requirement / Use Case / User Story / Acceptance Criteria")
doc.add_paragraph("- Administrator / Customer / Restaurant / Shipper: Approved Actors")
doc.add_paragraph("- COD / QR: Payment mechanisms")

# Fix font in tables
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    if r.font.size is None:
                        r.font.size = Pt(9)

doc.save(r'c:\Users\ACER\Downloads\Online-Delivery-Systems-Web\docs\ba-portfolio\03_Requirements\Product_Backlog_Online_Food_Delivery.docx')
print("DOCX created successfully.")
